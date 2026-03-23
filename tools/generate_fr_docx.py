"""
Generate NetPulse-DSIG-Standard-v0.5-FR.docx
D-SIG — Standard Signal Distillé v0.5 — Traduction française
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles globaux ──────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def quote(text):
    p = doc.add_paragraph(style='Quote')
    p.add_run(text).italic = True
    return p

def code_block(text):
    p = doc.add_paragraph(style='No Spacing')
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    # Rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()
    return table

def hr():
    doc.add_paragraph('─' * 72)

# ══════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════
title = doc.add_heading('D-SIG — Standard Signal Distillé v0.5', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Un plan architectural et un document de positionnement\npour la distillation des signaux de vitalité numériques').italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Domaine public — CC0 — Mars 2026\n')
meta.add_run('Implémentation de référence : NetPulse — github.com/dravitch/netpulse\n')
meta.add_run('Standard source : github.com/dravitch/dsig-standard')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. QU'EST-CE QUE D-SIG
# ══════════════════════════════════════════════════════════════════
h1('1. Qu\'est-ce que D-SIG')

para('D-SIG propose un cadre — pas une formule — pour produire, transmettre et consommer des signaux de vitalité unifiés dans les systèmes distribués.')
para('')
para('Il ne prétend pas calculer la vérité. Il propose une façon structurée de voir et comprendre différemment l\'état opérationnel : par convergence multi-perspectives, confiance cumulative, et distillation sémantiquement riche.')
para('')
quote('Comme TCP/IP ne garantit pas la livraison des paquets mais en maximise la probabilité via un protocole documenté, D-SIG ne calcule pas la vérité mais maximise la probabilité de produire une intelligence opérationnelle actionnable via une architecture de distillation documentée.')

hr()

# ══════════════════════════════════════════════════════════════════
# 2. LE SIGNAL — STRUCTURE MINIMALE
# ══════════════════════════════════════════════════════════════════
h1('2. Le Signal — Structure Minimale')

para('Chaque signal D-SIG est une Triple-Réduction : trois représentations obligatoires, atomiquement cohérentes, du même état, chacune ciblant un canal cognitif distinct.')

add_table(
    ['Composant', 'Forme', 'Signe (Peirce)', 'Canal cognitif'],
    [
        ['score', 'Entier 0–100', 'Index', 'Analytique — comparaison quantitative, consommation IA'],
        ['label', '4 états', 'Symbole', 'Sémantique — vocabulaire partagé tous niveaux d\'expertise'],
        ['color', '4 états', 'Icône', 'Limbique — réflexe pré-analytique, priorité visuelle immédiate'],
    ]
)

para('Plus deux champs temporels obligatoires :')
para('• trend — dérivée temporelle : STABLE | IMPROVING | DEGRADING | CRITICAL_FALL', bold=False)
para('• timestamp + ttl — validation de fraîcheur et anti-rejeu')

para('')
para('Règle 10 (nouveau v0.5) :', bold=True)
para('label est toujours f(score). trend est toujours f(d(score)/dt). Ils sont indépendants. Un label GOOD avec une tendance CRITICAL_FALL ne devient pas DEGRADED — il reste GOOD avec un signal d\'urgence séparé.')

h2('Table de correspondance')
add_table(
    ['Score', 'Label', 'Couleur', 'Hex'],
    [
        ['85–100', 'EXCELLENT', 'VERT', '#2ECC71'],
        ['60–84', 'GOOD', 'JAUNE', '#F4D03F'],
        ['35–59', 'DEGRADED', 'ORANGE', '#E67E22'],
        ['0–34', 'CRITICAL', 'ROUGE', '#E74C3C'],
    ]
)

hr()

# ══════════════════════════════════════════════════════════════════
# 3. CHAMPS REQUIS
# ══════════════════════════════════════════════════════════════════
h1('3. Champs Requis')

code_block('{\n  "dsig_version": "0.5",\n  "score": 82,\n  "label": "GOOD",\n  "color": "YELLOW",\n  "trend": "STABLE",\n  "timestamp": "2026-03-09T14:32:00Z",\n  "ttl": 900,\n  "source_id": "fp:a3:7c:12:...:ed25519",\n  "perspective": "LOCAL"\n}')

h2('Champs recommandés')
code_block('{\n  "source_sig": "base64(signature_Ed25519)",\n  "source_pub": "base64(empreinte_cle_publique)",\n  "trust_level": 1,\n  "baseline_cycles": 10,\n  "dimensions": {\n    "vital":      {"score": 1,  "ts": "...", "ttl": 300},\n    "local":      {"score": 10, "ts": "...", "ttl": 3600},\n    "internet":   {"score": 25, "ts": "...", "ttl": 3600},\n    "dns":        {"score": 15, "ts": "...", "ttl": 3600},\n    "throughput": {"score": 22, "ts": "...", "ttl": 43200},\n    "hub":        {"score": 10, "ts": "...", "ttl": 600}\n  },\n  "flags": []\n}')

para('Voir schema/dsig-signal.json pour le validateur JSON Schema complet.')

hr()

# ══════════════════════════════════════════════════════════════════
# 4. CINQ PRINCIPES ARCHITECTURAUX
# ══════════════════════════════════════════════════════════════════
h1('4. Cinq Principes Architecturaux')

add_table(
    ['#', 'Principe', 'Idée centrale', 'Origine'],
    [
        ['P1', 'Triple-Réduction Sémantique', 'score + label + color sont obligatoires et atomiquement cohérents', 'Sémiologie (Peirce, 1867)'],
        ['P2', 'Multi-Perspective Convergente', 'N perspectives indépendantes — convergence confirme, divergence diagnostique', 'Systèmes distribués'],
        ['P3', 'Autonomie de Production', 'Un nœud émet même sans réseau. Le silence est sémantiquement chargé.', 'Résilience Edge'],
        ['P4', 'Absorption Asymétrique du Bruit', 'Seules les tendances soutenues déclenchent des changements de label. Le silence est l\'état par défaut.', 'Traitement du signal'],
        ['P5', 'Diffusion Agnostique', 'Tout support de transport. Le récepteur n\'a pas besoin de connaître les sources sous-jacentes.', 'Stigmergie (Grassé, 1959)'],
    ]
)

hr()

# ══════════════════════════════════════════════════════════════════
# 5. DEUX PROPRIÉTÉS ÉMERGENTES
# ══════════════════════════════════════════════════════════════════
h1('5. Deux Propriétés Émergentes')

h2('Principe Prusik')
para('Chaque perspective légitime supplémentaire qui maintient son signal réel augmente la visibilité de la divergence par rapport aux signaux falsifiés. L\'attaque amplifie l\'alerte qu\'elle cherche à masquer. Tolérance byzantine : ⌊(N−1)/3⌋ perspectives corrompues tolérées avec N sources indépendantes.')

h2('Principe Phéromone')
para('Chaque cycle de convergence dépose de la confiance. Une rupture après une convergence soutenue (baseline_cycles élevé) est statistiquement plus significative qu\'une rupture dans un nouveau déploiement. La confiance s\'accumule ; elle s\'évapore aussi sans renouvellement (TTL).')

hr()

# ══════════════════════════════════════════════════════════════════
# 6. FORMULE DE DISTILLATION
# ══════════════════════════════════════════════════════════════════
h1('6. Formule de Distillation')

code_block('score = ( Σ wᵢ × normalize(mᵢ) ) × Π fail_fast_modifier(dⱼ) × precondition_gate(pₖ)')

para('• wᵢ — poids des dimensions. Somme = 1,0. Défini par l\'implémentation.')
para('• fail_fast_modifier — si une dimension critique = 0, le score s\'effondre vers CRITICAL (Règle 1)')
para('• precondition_gate (nouveau v0.5) — si la précondition échoue, le score est plafonné à ≤ 20, pas zéro (Règle 11)')

h2('Profil de référence : Nœud IT (NetPulse)')

add_table(
    ['Dimension', 'Poids', 'Type', 'Collecte'],
    [
        ['vital', 'PRÉCONDITION', 'precondition_gate', 'Chaque cycle'],
        ['local', '10 pts', 'graduel', 'Toutes les 6h'],
        ['internet', '25 pts', 'binaire', 'Toutes les 6h'],
        ['dns', '15 pts', 'binaire', 'Toutes les 6h'],
        ['throughput', '35 pts', 'graduel', 'Deux fois par jour'],
        ['hub', '15 pts', 'semi-binaire', 'Chaque cycle'],
    ]
)

para('Les implémentations DOIVENT documenter leur profil. Des profils standardisés (Infrastructure-Critique, IoT-Edge, Cloud-Native) sont prévus pour v1.0.')

hr()

# ══════════════════════════════════════════════════════════════════
# 7. RÉSUMÉ DES RÈGLES
# ══════════════════════════════════════════════════════════════════
h1('7. Résumé des Règles')

add_table(
    ['Règle', 'Nom', 'Version'],
    [
        ['1', 'Fail-Fast sur la Dimension Vitale', 'v0.1'],
        ['2', 'Lissage Temporel du Bruit', 'v0.1'],
        ['3', 'La Divergence de Perspective est Diagnostique', 'v0.1'],
        ['4', 'Autonomie de Production', 'v0.1'],
        ['5', 'Immuabilité de la Triple-Réduction', 'v0.1'],
        ['6', 'Indépendance des Sources Vérifiable (Byzantin : ⌊(N−1)/3⌋)', 'v0.3'],
        ['7', 'Rejet du Signal Contradictoire (incl. saut baseline_cycles)', 'v0.3'],
        ['8', 'Signal STALE / Anti-Rejeu (tolérance horloge ±30s)', 'v0.3'],
        ['9', 'Historique de Convergence Calculé par le Récepteur', 'v0.4, révisé v0.5'],
        ['10', 'Indépendance Label/Tendance (label = f(score) uniquement)', 'nouveau v0.5'],
        ['11', 'Dimension Précondition (plafond ≤ 20, pas zéro)', 'nouveau v0.5'],
    ]
)

hr()

# ══════════════════════════════════════════════════════════════════
# 8. MODÈLE DE CONFIANCE PRODUCTEUR (D-SIG-PROD)
# ══════════════════════════════════════════════════════════════════
h1('8. Modèle de Confiance Producteur (D-SIG-PROD)')

add_table(
    ['ID', 'Nom', 'Résumé'],
    [
        ['PROD-01', 'Identité Stable', 'source_id = clé cryptographique, pas adresse réseau'],
        ['PROD-02', 'Sémantique de l\'Absence', 'Le silence est un signal, pas un vide'],
        ['PROD-03', 'Co-Habitation Supposée', 'Contrôle et observation partagent le même réseau par défaut'],
        ['PROD-04', 'Authenticité du Signal Distincte du Canal', 'source_sig ≠ authentification SSH'],
        ['PROD-05', 'Ancre Unique — Invariant Systémique', 'La robustesse D-SIG est bornée par la sécurité du mécanisme d\'enrôlement'],
    ]
)

hr()

# ══════════════════════════════════════════════════════════════════
# 9. LIMITES DOCUMENTÉES
# ══════════════════════════════════════════════════════════════════
h1('9. Limites Documentées')

para('1. Risque de sur-distillation — des poids mal calibrés créent une fausse confiance')
para('2. Effet observateur — la sonde perturbe ce qu\'elle mesure (minimisé, non résolu)')
para('3. Problème de l\'oracle — un producteur peut signer des signaux valides avec des données fabriquées ; atténué par la convergence multi-perspectives, non résolu')
para('4. Ancre unique — toute résistance byzantine échoue si le mécanisme d\'enrôlement est compromis')
para('5. Interopérabilité des pondérations — deux déploiements D-SIG-compatibles avec des profils différents produisent des scores non comparables')

hr()

# ══════════════════════════════════════════════════════════════════
# 10. GÉNÉALOGIE INTELLECTUELLE
# ══════════════════════════════════════════════════════════════════
h1('10. Généalogie Intellectuelle')

para('D-SIG applique l\'architecture des indicateurs composites du trading algorithmique à l\'observabilité IT :')

add_table(
    ['QAAF / Trading', 'D-SIG', 'Rôle'],
    [
        ['Métriques brutes prix/volume', 'Métriques réseau brutes', 'Couche d\'entrée'],
        ['Score composite (0–100)', 'Score distillé (0–100)', 'Vitalité normalisée'],
        ['Signal de tendance MACD, RSI', 'trend : IMPROVING / DEGRADING / CRITICAL_FALL', 'Dérivée temporelle'],
        ['Déclencheur stop-loss', 'Modificateur fail-fast', 'Seuil critique'],
        ['Période de référence RSI', 'baseline_cycles', 'Proxy de confiance accumulée'],
        ['ACHETER / CONSERVER / VENDRE', 'EXCELLENT / GOOD / DEGRADED / CRITICAL', 'Sortie sémantique actionnable'],
    ]
)

hr()

# ══════════════════════════════════════════════════════════════════
# 11. GOUVERNANCE
# ══════════════════════════════════════════════════════════════════
h1('11. Gouvernance')

para('Licence : CC0 — Domaine Public. Aucune restriction. Aucune autorité de certification.', bold=True)
para('')
para('Art antérieur : horodatage du commit GitHub + soumission arXiv cs.NI')
para('Implémentation de référence : NetPulse — copyright auteur, séparé de ce standard')
para('Compatibilité : implémenter les champs requis + documenter son profil de pondération = D-SIG compatible')

doc.add_paragraph()
quote('D-SIG est domaine public (CC0). Publié sur GitHub et arXiv pour établir l\'antériorité. Personne ne peut le breveter — y compris son auteur.')

hr()

# ══════════════════════════════════════════════════════════════════
# PIED DE PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('D-SIG Standard v0.5 · Domaine Public CC0 · Mars 2026\n').bold = True
footer.add_run('Implémentation de référence : NetPulse v1.7+ · github.com/dravitch/netpulse\n')
footer.add_run('Standard source : github.com/dravitch/dsig-standard')

# ── Sauvegarde ──────────────────────────────────────────────────
output_path = '/home/netpulser/dsig-standard/NetPulse-DSIG-Standard-v0.5-FR.docx'
doc.save(output_path)
print(f'OK: {output_path}')
